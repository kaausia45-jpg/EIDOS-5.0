# -*- coding: utf-8 -*-
"""[2026-06-03 ②] AIRA 문서·사진·영상 "열어서 설명" 핵심 로직.

순수 데이터/생성 로직(EIDOS GUI 무의존·LLM 주입형·테스트 가능). GUI(AiraDocDialog·
AiraImageDialog·AiraVideoDialog·레이저 포인터)는 eidos_chat_gui.py 에 둔다.

모드별 정의(사용자 결정):
- 문서(doc)  : AIRA 가 내용을 **직접 생성** → 섹션별 음성 설명 + 레이저.
- 사진(photo): **기존 파일** 열기 → 비전 LLM 이 보고 설명(음성) → 영역 레이저.
- 영상(video): **기존 파일** 열기 → 재생 + AIRA 소개/설명 음성(생성 X).
- PPT        : P4(eidos_aira_deck)에서 완료.

설계:
- generate_doc_async: 주제 → LLM → 섹션 구조 JSON. 실패 시 안전한 1섹션 폴백.
- describe_image_async: 비전 LLM(image_input) → 설명 + 짚을 영역(regions·0~1 정규화).
- video_intro: 파일 메타 + 제목 기반 소개 narration(동기·LLM 선택).
- classify_target: 경로/주제 → "doc|image|video|unknown" 분기.
- doc_to_docx: python-docx 로 .docx 내보내기(있을 때만).
"""
from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, field

# 표정 키 힌트 — assets/expr_*.png 에 실제 존재하는 것들(생성기가 이 안에서 고르게 유도).
# GUI 가 실제 스캔 목록으로 최종 검증하므로 여기 누락돼도 안전.
_EXPR_HINTS = (
    "information", "point", "teaching", "excited", "happy",
    "confidence", "question", "thinking", "neutral", "ending", "surprised",
)

# 파일 타입 판별용 확장자
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tif", ".tiff"}
_VIDEO_EXTS = {".mp4", ".avi", ".mkv", ".mov", ".wmv", ".webm", ".flv", ".m4v", ".mpg", ".mpeg"}
_DOC_EXTS = {".txt", ".md", ".markdown", ".doc", ".docx", ".rtf", ".pdf"}


# ── 데이터 타입 ─────────────────────────────────────────────────────────────
@dataclass
class Section:
    heading: str = ""
    body: str = ""                                 # 본문(여러 문장)
    narration: str = ""                            # AIRA 가 말할 대사
    expr: str = "information"                       # 표정 키 힌트

    def to_dict(self) -> dict:
        return {
            "heading": self.heading,
            "body": self.body,
            "narration": self.narration,
            "expr": self.expr,
        }

    @staticmethod
    def from_dict(d: dict) -> "Section":
        d = d or {}
        body = d.get("body", "")
        if isinstance(body, (list, tuple)):
            body = "\n".join(str(x).strip() for x in body if str(x).strip())
        return Section(
            heading=str(d.get("heading", "") or "").strip(),
            body=str(body or "").strip(),
            narration=str(d.get("narration", "") or "").strip(),
            expr=str(d.get("expr", "information") or "information").strip().lower(),
        )


@dataclass
class Doc:
    title: str = ""
    sections: list = field(default_factory=list)   # list[Section]
    topic: str = ""

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "topic": self.topic,
            "sections": [s.to_dict() for s in self.sections],
        }

    @staticmethod
    def from_dict(d: dict) -> "Doc":
        d = d or {}
        return Doc(
            title=str(d.get("title", "") or "").strip(),
            topic=str(d.get("topic", "") or "").strip(),
            sections=[Section.from_dict(s) for s in (d.get("sections") or [])],
        )

    def to_markdown(self) -> str:
        lines = [f"# {self.title or self.topic or '문서'}", ""]
        for s in self.sections:
            if s.heading:
                lines.append(f"## {s.heading}")
            if s.body:
                lines.append(s.body)
            lines.append("")
        return "\n".join(lines).strip() + "\n"


# ── 파일 타입 판별 ──────────────────────────────────────────────────────────
def classify_target(path_or_topic: str) -> str:
    """경로면 확장자로 image/video/doc 판별, 실제 파일 아니면 unknown(=주제→문서 생성).

    반환: "image" | "video" | "doc" | "unknown".
    """
    s = (path_or_topic or "").strip().strip('"').strip("'")
    if not s:
        return "unknown"
    ext = os.path.splitext(s)[1].lower()
    if ext in _IMAGE_EXTS:
        return "image"
    if ext in _VIDEO_EXTS:
        return "video"
    if ext in _DOC_EXTS:
        return "doc"
    # 확장자 없음/모르는 확장자 → 실제 존재 파일이면 doc 시도, 아니면 unknown(주제)
    if os.path.isfile(s):
        return "doc"
    return "unknown"


# ── JSON 파싱(실 LLM 강건성·deck 모듈과 동일 전략) ──────────────────────────
def _strip_fences(text: str) -> str:
    t = (text or "").strip()
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z0-9]*\s*", "", t)
        if t.endswith("```"):
            t = t[:-3]
    return t.strip()


def _extract_balanced(text: str, open_ch: str, close_ch: str) -> str:
    """문자열 리터럴 안의 괄호는 무시하고 균형 잡힌 첫 블록 추출."""
    s = text
    start = s.find(open_ch)
    if start < 0:
        return ""
    depth = 0
    in_str = False
    esc = False
    quote = ""
    for i in range(start, len(s)):
        c = s[i]
        if in_str:
            if esc:
                esc = False
            elif c == "\\":
                esc = True
            elif c == quote:
                in_str = False
            continue
        if c in ('"', "'"):
            in_str = True
            quote = c
        elif c == open_ch:
            depth += 1
        elif c == close_ch:
            depth -= 1
            if depth == 0:
                return s[start:i + 1]
    return ""


def _safe_json(text: str):
    t = _strip_fences(text)
    try:
        return json.loads(t)
    except Exception:
        pass
    block = _extract_balanced(t, "{", "}")
    if block:
        try:
            return json.loads(block)
        except Exception:
            pass
    return None


# ── 문서 생성(직접 생성) ────────────────────────────────────────────────────
async def generate_doc_async(
    topic: str,
    *,
    llm_async,
    n_sections: int = 5,
    audience: str = "",
    timeout_sec: float = 60.0,
) -> Doc:
    """주제 → AIRA 가 직접 쓴 문서(LLM). 실패 시 1섹션 안전 폴백.

    llm_async: get_llm_response_async 호환(prompt, max_tokens=..., use_cache=...).
    """
    topic = (topic or "").strip()
    if not topic:
        return Doc(title="문서", topic=topic, sections=[
            Section(heading="문서 주제가 비어 있어요",
                    body="어떤 문서를 만들어 드릴까요?",
                    narration="어떤 문서를 만들어 드릴까요? 주제를 알려주시면 제가 작성해서 설명해 드릴게요.",
                    expr="question")])

    n = max(3, min(int(n_sections or 5), 10))
    aud = f"\n독자: {audience.strip()}" if (audience or "").strip() else ""
    prompt = (
        "너는 ENFP 안내자 AIRA 야. 아래 주제로 짧은 문서를 JSON 으로 작성해.\n"
        f"주제: {topic}{aud}\n"
        f"- 섹션 {n}개 내외. 각 섹션은 heading(짧게), "
        "body(2~4문장의 실제 내용·평문), "
        "narration(AIRA 가 그 섹션을 짚으며 말할 대사·친근한 존댓말 해요체·반말 금지·1~3문장), "
        "expr(표정 키 1개).\n"
        f"- expr 는 다음 중 하나: {', '.join(_EXPR_HINTS)}.\n"
        "JSON 형식만 출력(설명 금지):\n"
        '{"title":"문서 제목","sections":[{"heading":"...","body":"...",'
        '"narration":"...","expr":"information"}]}'
    )
    raw = ""
    try:
        import asyncio
        raw = await asyncio.wait_for(
            llm_async(prompt, max_tokens=2400, use_cache=False), timeout=timeout_sec)
    except Exception as e:
        raw = ""
        print(f"[aira explain] doc 생성 실패: {e}")

    data = _safe_json(raw) if raw else None
    if not isinstance(data, dict) or not data.get("sections"):
        return Doc(
            title=topic[:60], topic=topic,
            sections=[Section(
                heading=topic[:60],
                body="문서 자동 생성에 실패했어요. 주제를 더 구체적으로 알려주세요.",
                narration=f"{topic} 문서를 준비하려 했는데 생성이 잘 안 됐어요. 주제를 조금 더 구체적으로 알려주시겠어요?",
                expr="concerned")])

    doc = Doc.from_dict({**data, "topic": topic})
    if not doc.title:
        doc.title = topic[:60]
    doc.sections = [s for s in doc.sections if (s.heading or s.body or s.narration)]
    if not doc.sections:
        doc.sections = [Section(heading=topic[:60], body=topic[:120],
                                narration=f"{topic} 에 대해 설명해 드릴게요.", expr="information")]
    return doc


# ── 사진 설명(생성 X·기존 파일·비전 LLM) ────────────────────────────────────
def _clamp01(v, default=0.5):
    try:
        f = float(v)
    except Exception:
        return default
    return max(0.0, min(1.0, f))


def _load_image_bytes(image_path: str):
    """이미지 파일 → bytes. 실패 시 None."""
    try:
        with open(image_path, "rb") as f:
            return f.read()
    except Exception as e:
        print(f"[aira explain] 이미지 로드 실패: {e}")
        return None


async def describe_image_async(
    image_path: str,
    *,
    llm_async,
    timeout_sec: float = 60.0,
) -> dict:
    """기존 사진을 비전 LLM 으로 보고 설명 + 짚을 영역 힌트.

    llm_async: get_llm_response_async 호환(image_input=bytes 멀티모달 지원).
    반환: {"description": str, "narration": str,
           "regions": [{"label": str, "x": 0~1, "y": 0~1}], "expr": str}.
    """
    image_path = (image_path or "").strip().strip('"').strip("'")
    if not image_path or not os.path.isfile(image_path):
        return {"description": "", "narration": "사진 파일을 찾을 수 없어요. 경로를 확인해 주시겠어요?",
                "regions": [], "expr": "concerned"}
    img = _load_image_bytes(image_path)
    if img is None:
        return {"description": "", "narration": "사진을 여는 데 실패했어요.",
                "regions": [], "expr": "concerned"}

    prompt = (
        "너는 ENFP 안내자 AIRA 야. 이 사진을 보고 옆에서 짚어주며 설명할 거야.\n"
        "친근한 존댓말 해요체로 말해(반말 금지).\n"
        "- description: 사진 전체를 2~4문장으로 설명.\n"
        "- narration: AIRA 가 음성으로 말할 대사(1~3문장·존댓말).\n"
        "- regions: 짚어줄 핵심 지점 1~4개. 각 {label(짧은 이름), x, y} — "
        "x,y 는 사진 좌상단(0,0)~우하단(1,1) 기준 0~1 정규화 좌표.\n"
        f"- expr: 다음 중 하나({', '.join(_EXPR_HINTS)}).\n"
        "JSON 만 출력(설명 금지):\n"
        '{"description":"...","narration":"...",'
        '"regions":[{"label":"...","x":0.5,"y":0.4}],"expr":"point"}'
    )
    raw = ""
    try:
        import asyncio
        raw = await asyncio.wait_for(
            llm_async(prompt, image_input=img), timeout=timeout_sec)
    except Exception as e:
        print(f"[aira explain] 사진 설명 실패: {e}")
        raw = ""

    data = _safe_json(raw) if raw else None
    if not isinstance(data, dict):
        return {"description": "", "narration": "지금은 사진을 설명하기 어려웠어요. 다시 시도해 주시겠어요?",
                "regions": [], "expr": "concerned"}

    regions = []
    for r in (data.get("regions") or []):
        if not isinstance(r, dict):
            continue
        regions.append({
            "label": str(r.get("label", "") or "").strip()[:40],
            "x": _clamp01(r.get("x", 0.5)),
            "y": _clamp01(r.get("y", 0.5)),
        })
    expr = str(data.get("expr", "point") or "point").strip().lower()
    return {
        "description": str(data.get("description", "") or "").strip(),
        "narration": str(data.get("narration", "") or "").strip()
                     or str(data.get("description", "") or "").strip(),
        "regions": regions[:4],
        "expr": expr,
    }


# ── 영상 소개(생성 X·기존 파일·메타+제목) ───────────────────────────────────
def _human_duration(sec) -> str:
    try:
        sec = int(float(sec))
    except Exception:
        return ""
    if sec <= 0:
        return ""
    m, s = divmod(sec, 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h}시간 {m}분 {s}초"
    if m:
        return f"{m}분 {s}초"
    return f"{s}초"


def video_meta(video_path: str) -> dict:
    """영상 파일 메타(이름·크기·확장자). 길이/해상도는 GUI(QMediaPlayer)에서 채울 수 있음."""
    out = {"path": video_path, "name": "", "ext": "", "size_mb": 0.0,
           "duration_sec": 0, "exists": False}
    try:
        if video_path and os.path.isfile(video_path):
            out["exists"] = True
            out["name"] = os.path.basename(video_path)
            out["ext"] = os.path.splitext(video_path)[1].lower()
            out["size_mb"] = round(os.path.getsize(video_path) / (1024 * 1024), 1)
    except Exception:
        pass
    return out


def video_intro(video_path: str, *, duration_sec: int = 0, llm_async=None) -> dict:
    """기존 영상 소개 — 파일 메타 + 제목 기반 소개 narration(생성 X).

    llm_async 는 받지만 동기 경로(메타 기반)만 사용 — 외부 의존·블로킹 회피.
    반환: {"title": str, "narration": str, "meta": dict, "expr": str}.
    """
    meta = video_meta(video_path)
    if duration_sec:
        meta["duration_sec"] = int(duration_sec)
    if not meta["exists"]:
        return {"title": "", "meta": meta, "expr": "concerned",
                "narration": "영상 파일을 찾을 수 없어요. 경로를 확인해 주시겠어요?"}

    name = os.path.splitext(meta["name"])[0]
    dur = _human_duration(meta.get("duration_sec"))
    bits = [f"'{name}' 영상을 열었어요."]
    if dur:
        bits.append(f"길이는 {dur} 예요.")
    bits.append("지금부터 함께 보면서 설명해 드릴게요.")
    narration = " ".join(bits)
    return {"title": name, "meta": meta, "expr": "excited", "narration": narration}


# ── docx 내보내기(선택·python-docx 있을 때만) ───────────────────────────────
def docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def doc_to_docx(doc: Doc, path: str) -> str:
    """Doc → .docx 파일. python-docx 필요. 성공 시 path 반환, 실패 시 예외."""
    from docx import Document

    d = Document()
    d.add_heading(doc.title or (doc.topic or "문서")[:60], level=0)
    for s in doc.sections:
        if s.heading:
            d.add_heading(s.heading, level=1)
        if s.body:
            d.add_paragraph(s.body)
    d.save(path)
    return path
